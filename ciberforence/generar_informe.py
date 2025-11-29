#!/usr/bin/env python3

import json
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt


def generar_informe(caso, cliente, resultados_file, contrato):
    try:
        with open(resultados_file, encoding="utf-8") as f:
            datos = json.load(f)
    except FileNotFoundError:
        print(
            f"El archivo de resultados '{resultados_file}' no se encontró. Por favor, verifica el nombre y la ubicación."
        )
        sys.exit(1)
    except json.JSONDecodeError:
        print(
            f"El archivo '{resultados_file}' no contiene datos válidos. Asegúrate de que sea un archivo JSON correcto."
        )
        sys.exit(1)

    if not isinstance(datos, list):
        print(
            "El archivo de resultados no contiene una lista de datos. Por favor, revisa el formato del archivo."
        )
        sys.exit(1)

    doc = Document()
    estilo = doc.styles["Normal"]
    fuente = estilo.font
    fuente.name = "Arial"
    fuente.size = Pt(11)

    doc.add_heading("INFORME PERICIAL TÉCNICO - DELITOS INFORMÁTICOS", 0)

    doc.add_heading("1. Datos del caso", level=1)
    doc.add_paragraph(f"• Número de caso: {caso}")
    doc.add_paragraph(f"• Cliente: {cliente}")
    doc.add_paragraph(f"• Contrato de autorización: {contrato}")
    doc.add_paragraph(f"• Fecha de elaboración: {datetime.now().strftime('%d/%m/%Y')}")

    doc.add_heading("2. Metodología", level=1)
    doc.add_paragraph(
        "Para este informe se realizó un análisis pasivo de inteligencia de amenazas sobre las direcciones IP reportadas como origen de accesos no autorizados. "
        "No se efectuó ningún escaneo activo ni interacción directa con sistemas externos, más allá de consultas a fuentes públicas como IPinfo, AbuseIPDB y WHOIS. "
        "El objetivo fue comprender la naturaleza técnica y el comportamiento de cada IP, para distinguir entre posibles atacantes y dispositivos víctimas de malware."
    )

    doc.add_heading("3. Hallazgos por IP", level=1)
    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = "Table Grid"
    encabezados = tabla.rows[0].cells
    encabezados[0].text = "IP"
    encabezados[1].text = "País / ISP"
    encabezados[2].text = "Reputación (AbuseIPDB)"
    encabezados[3].text = "Clasificación"

    for resultado in datos:
        if "error" not in resultado:
            fila = tabla.add_row().cells
            fila[0].text = resultado.get("ip", "")
            fila[1].text = f"{resultado.get('pais', '')} / {resultado.get('isp', '')}"
            fila[2].text = f"{resultado.get('abuse_score', '')}%"
            fila[3].text = resultado.get("clasificacion", "")

    doc.add_heading("4. Conclusiones", level=1)
    doc.add_paragraph(
        "El análisis permitió identificar direcciones IP asociadas a posibles actores externos con intención delictiva, así como dispositivos domésticos en Ecuador que podrían estar comprometidos por malware. "
        "Se recomienda presentar denuncia únicamente contra las IPs clasificadas como 'atacante real' y coordinar con los proveedores de internet locales para ayudar a limpiar los nodos infectados."
    )

    doc.add_heading("5. Marco legal", level=1)
    doc.add_paragraph(
        f"Este informe se emite en el marco del Artículo 221 del Código Orgánico Integral Penal (COIP) de la República del Ecuador, que sanciona el acceso no consentido a sistemas informáticos. "
        f"La investigación se realizó bajo el contrato N°{contrato}, con la autorización expresa del cliente."
    )

    nombre_doc = f"informe_fiscalia_{caso.replace('/', '_')}.docx"
    doc.save(nombre_doc)
    print(f"Informe generado exitosamente: {nombre_doc}")


if __name__ == "__main__":
    if len(sys.argv) != 5:
        print(
            "Uso correcto: python 02_generar_informe.py <caso> <cliente> <resultados.json> <contrato>"
        )
        sys.exit(1)
    generar_informe(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4])
