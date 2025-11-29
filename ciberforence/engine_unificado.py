#!/usr/bin/env python3
"""
ENGINE UNIFICADO FORENSE PASIVO (DedSec Edition)
------------------------------------------------
Este motor centraliza:
- Localización OSINT (IPinfo)
- Reputación AbuseIPDB
- Clasificación avanzada (VPS, Internacional, Botnet, Hogar)
- Diccionario local (rangos, ASN, listas negras)
- Generación de informe pericial en DOCX

Modo de uso:
    python3 engine_unificado.py ips.txt --caso 2025-045 --cliente "Empresa S.A." --contrato "789" --abuse CLAVE --token TOKEN

Requiere:
- diccionario_ips.json (en el directorio de ejecución)
- asn_ecuador.json (en el directorio de ejecución)
- Paquetes externos: requests, python-docx

Este archivo es autocontenible en cuanto a lógica. Los datos externos (JSON) y librerías se
deben proveer por separado.
"""

import argparse
import json
import os
import socket
import sys
from datetime import datetime

import requests

try:
    from docx import Document
    from docx.shared import Pt

    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False
    Document = None
    Pt = None


# =========================
# Configuración / Constantes
# =========================
DICCIONARIO_PATH = "diccionario_ips.json"
ASN_ECUADOR_PATH = "asn_ecuador.json"


# =========================
# Módulo: Lector Diccionario
# =========================
def cargar_diccionario(path=DICCIONARIO_PATH):
    """
    Carga el diccionario local de IPs desde un archivo JSON.
    Si el archivo no existe, retorna un diccionario vacío.
    """
    if not os.path.exists(path):
        return {}
    with open(path, "r", encoding="utf-8") as f:
        try:
            return json.load(f)
        except Exception:
            return {}


def ip_en_diccionario(ip, diccionario):
    """
    Consulta si una IP está en el diccionario y retorna su categoría o información asociada.
    Si no está, retorna None.
    """
    return diccionario.get(ip, None)


# =========================
# Módulo: Localización OSINT
# =========================
def ipinfo_lookup(ip, token=None):
    """
    Consulta IPinfo para obtener datos JSON crudos de una IP.
    """
    url = f"https://ipinfo.io/{ip}/json"
    if token:
        url += f"?token={token}"
    try:
        res = requests.get(url, timeout=8)
        return res.json() if res.status_code == 200 else {}
    except Exception as e:
        return {"error": str(e)}


def localizar_ip(ip, token=None):
    """
    Geolocalización resumida de una IP usando IPinfo.
    Devuelve un diccionario con país, región, ciudad, organización y coordenadas.
    """
    data = ipinfo_lookup(ip, token)
    if data and "error" not in data:
        return {
            "ip": ip,
            "pais": data.get("country", "N/A"),
            "region": data.get("region", "N/A"),
            "ciudad": data.get("city", "N/A"),
            "org": data.get("org", "N/A"),
            "loc": data.get("loc", "N/A"),
        }
    return {
        "ip": ip,
        "pais": "N/A",
        "region": "N/A",
        "ciudad": "N/A",
        "org": "N/A",
        "loc": "N/A",
    }


# =========================
# Módulo: Reputación / Clasificador
# =========================
def abuseipdb_lookup(ip, api_key):
    """
    Consulta AbuseIPDB para obtener reputación de una IP.
    """
    url = "https://api.abuseipdb.com/api/v2/check"
    headers = {"Key": api_key, "Accept": "application/json"}
    params = {"ipAddress": ip, "maxAgeInDays": "90"}
    try:
        res = requests.get(url, headers=headers, params=params, timeout=8)
        data = res.json()
        return data.get("data", {}) if res.status_code == 200 else {}
    except Exception as e:
        return {"error": str(e)}


def _cargar_asn_ecuador(path=ASN_ECUADOR_PATH):
    """
    Carga mapa ASN de Ecuador desde JSON. Si no existe o es inválido, devuelve {}.
    Formato esperado: { "AS12345": "Nombre ISP", ... }
    """
    if not os.path.exists(path):
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def clasificar_ip(ip, abuse_key, ipinfo_token=None):
    """
    Clasifica una IP en función de:
    - País (internacional vs Ecuador)
    - Organización/ASN (indicios de datacenter/VPS)
    - Puntuación de reputación (AbuseIPDB)

    Retorna un dict con atributos y una conclusión textual.
    """
    try:
        socket.inet_aton(ip)
    except Exception:
        return {"ip": ip, "error": "IP inválida"}

    ipinfo = ipinfo_lookup(ip, ipinfo_token)
    abuse = abuseipdb_lookup(ip, abuse_key)
    asn_ec_map = _cargar_asn_ecuador()

    org = ipinfo.get("org", "N/A")
    country = ipinfo.get("country", "N/A")
    asn_raw = org.split()[0] if org != "N/A" and org.split() else "N/A"

    is_ec = country == "EC"
    isp_ec = asn_ec_map.get(asn_raw, None) if is_ec else None

    is_datacenter = False
    if org != "N/A":
        org_low = org.lower()
        datacenter_keywords = [
            "cloud",
            "hosting",
            "vps",
            "digitalocean",
            "aws",
            "hetzner",
            "linode",
            "netlife",
            "fibra",
            "datacenter",
        ]
        if any(kw in org_low for kw in datacenter_keywords):
            is_datacenter = True

    abuse_score = int(abuse.get("abuseConfidenceScore", 0))

    if not is_ec:
        conclusion = "Atacante internacional"
    elif is_datacenter:
        conclusion = "VPS en Ecuador (posible atacante real)"
    elif abuse_score > 60:
        conclusion = "IP con alta reputación maliciosa (posible botnet)"
    else:
        conclusion = "Dispositivo doméstico infectado (víctima)"

    return {
        "ip": ip,
        "pais": country,
        "isp": org,
        "asn": asn_raw,
        "isp_ecuador": isp_ec,
        "es_datacenter": is_datacenter,
        "abuse_score": abuse_score,
        "clasificacion": conclusion,
    }


def generar_informe(caso, cliente, resultados_file, contrato):
    """

    Genera informe pericial DOCX a partir de resultados en JSON.

    Espera que 'resultados_file' contenga una lista de dicts producidos por analizar_ips().

    """

    if not DOCX_AVAILABLE:
        print(
            "python-docx no está instalado. Se omite la generación de informe DOCX. Los resultados JSON han sido generados."
        )
        return
    try:
        with open(resultados_file, encoding="utf-8") as f:
            datos = json.load(f)

    except FileNotFoundError:
        print(
            f"El archivo de resultados '{resultados_file}' no se encontró. Por favor, verifica el nombre y la ubicación."
        )

        sys.exit(1)

    except json.JSONDecodeError:
        print(
            f"El archivo '{resultados_file}' no contiene datos válidos. Asegúrate de que sea un archivo JSON correcto."
        )

        sys.exit(1)

    if not isinstance(datos, list):
        print(
            "El archivo de resultados no contiene una lista de datos. Por favor, revisa el formato del archivo."
        )
        sys.exit(1)

    doc = Document()
    estilo = doc.styles["Normal"]
    fuente = estilo.font
    fuente.name = "Arial"
    fuente.size = Pt(11)

    doc.add_heading("INFORME PERICIAL TÉCNICO - DELITOS INFORMÁTICOS", 0)

    doc.add_heading("1. Datos del caso", level=1)
    doc.add_paragraph(f"• Número de caso: {caso}")
    doc.add_paragraph(f"• Cliente: {cliente}")
    doc.add_paragraph(f"• Contrato de autorización: {contrato}")
    doc.add_paragraph(f"• Fecha de elaboración: {datetime.now().strftime('%d/%m/%Y')}")

    doc.add_heading("2. Metodología", level=1)
    doc.add_paragraph(
        "Para este informe se realizó un análisis pasivo de inteligencia de amenazas sobre las direcciones IP reportadas como origen de accesos no autorizados. "
        "No se efectuó ningún escaneo activo ni interacción directa con sistemas externos, más allá de consultas a fuentes públicas como IPinfo, AbuseIPDB y WHOIS. "
        "El objetivo fue comprender la naturaleza técnica y el comportamiento de cada IP, para distinguir entre posibles atacantes y dispositivos víctimas de malware."
    )

    doc.add_heading("3. Hallazgos por IP", level=1)
    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = "Table Grid"
    encabezados = tabla.rows[0].cells
    encabezados[0].text = "IP"
    encabezados[1].text = "País / ISP"
    encabezados[2].text = "Reputación (AbuseIPDB)"
    encabezados[3].text = "Clasificación"

    # Esta tabla espera los datos normalizados del clasificador
    for resultado in datos:
        if "error" not in resultado:
            fila = tabla.add_row().cells
            fila[0].text = resultado.get("ip", "")
            fila[1].text = f"{resultado.get('pais', '')} / {resultado.get('isp', '')}"
            fila[2].text = f"{resultado.get('abuse_score', '')}%"
            fila[3].text = resultado.get("clasificacion", "")

    doc.add_heading("4. Conclusiones", level=1)
    doc.add_paragraph(
        "El análisis permitió identificar direcciones IP asociadas a posibles actores externos con intención delictiva, así como dispositivos domésticos en Ecuador que podrían estar comprometidos por malware. "
        "Se recomienda presentar denuncia únicamente contra las IPs clasificadas como 'atacante real' y coordinar con los proveedores de internet locales para ayudar a limpiar los nodos infectados."
    )

    doc.add_heading("5. Marco legal", level=1)
    doc.add_paragraph(
        f"Este informe se emite en el marco del Artículo 221 del Código Orgánico Integral Penal (COIP) de la República del Ecuador, que sanciona el acceso no consentido a sistemas informáticos. "
        f"La investigación se realizó bajo el contrato N°{contrato}, con la autorización expresa del cliente."
    )

    nombre_doc = f"informe_fiscalia_{caso.replace('/', '_')}.docx"
    doc.save(nombre_doc)
    print(f"Informe generado exitosamente: {nombre_doc}")


# =========================
# Motor Unificado
# =========================
def analizar_ips(ips, abuse_key, ipinfo_token=None, diccionario_path=DICCIONARIO_PATH):
    """
    Orquesta el análisis de una lista de IPs:
    - Geolocalización vía IPinfo
    - Reputación AbuseIPDB
    - Clasificación local según diccionario
    - Fusiona resultados y los guarda en resultado_unificado.json
    """
    dic = cargar_diccionario(diccionario_path)
    resultados = []

    for ip in ips:
        ip = ip.strip()
        if not ip:
            continue

        print(f"\n🔍 Analizando {ip}…")

        # OSINT: Geolocalización
        localizacion = localizar_ip(ip, ipinfo_token)

        # Reputación y clasificación
        rep = clasificar_ip(ip, abuse_key, ipinfo_token)

        # Clasificación local forense (diccionario)
        categoria_local = ip_en_diccionario(ip, dic)

        fusion = {
            "ip": ip,
            "localizacion": localizacion,
            "clasificacion_forense": rep,
            "categoria_local": categoria_local,
            "timestamp": datetime.utcnow().isoformat(),
        }

        resultados.append(fusion)

    # Para el informe, también generamos un archivo plano con la salida del clasificador
    resultados_clasificador = [
        r["clasificacion_forense"]
        for r in resultados
        if "error" not in r.get("clasificacion_forense", {})
    ]

    with open("resultado_unificado.json", "w", encoding="utf-8") as f:
        json.dump(resultados, f, indent=2, ensure_ascii=False)

    with open("resultados_ips.json", "w", encoding="utf-8") as f:
        json.dump(resultados_clasificador, f, indent=2, ensure_ascii=False)

    print("Archivos generados: resultado_unificado.json, resultados_ips.json")
    return resultados


def _leer_ips_desde_archivo(path):
    if not os.path.exists(path):
        print(f"No existe el archivo {path}")
        sys.exit(1)
    with open(path, "r", encoding="utf-8") as f:
        return [line.strip() for line in f.readlines() if line.strip()]


def main():
    parser = argparse.ArgumentParser(
        description="Motor unificado de análisis pasivo de IPs"
    )
    parser.add_argument("archivo_ips", help="Archivo con lista de IPs")
    parser.add_argument("--abuse", required=True, help="API Key de AbuseIPDB")
    parser.add_argument("--token", required=False, help="Token de IPinfo")
    parser.add_argument("--caso", required=True, help="Número de caso para informe")
    parser.add_argument(
        "--cliente", required=True, help="Cliente o entidad solicitante"
    )
    parser.add_argument("--contrato", required=True, help="Contrato/autorización")
    parser.add_argument(
        "--dic",
        required=False,
        default=DICCIONARIO_PATH,
        help="Ruta al diccionario local JSON",
    )

    args = parser.parse_args()

    ips = _leer_ips_desde_archivo(args.archivo_ips)
    print(f"Iniciando análisis de {len(ips)} IPs")

    analizar_ips(ips, args.abuse, args.token, args.dic)

    print("Generando informe DOCX…")
    # El informe usa la salida plana del clasificador: resultados_ips.json
    generar_informe(
        args.caso,
        args.cliente,
        "resultados_ips.json",
        args.contrato,
    )

    print("Proceso finalizado. Informe listo.")


if __name__ == "__main__":
    main()
